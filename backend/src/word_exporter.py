"""
Word 导出模块

将剧本 JSON 导出为 Word 文档，只包含对话和镜头描述，不含技术字段。
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def _set_chinese_font(run, font_name: str = "宋体") -> None:
    """设置中文字体（同时设置东亚字体）。"""
    run.font.name = font_name
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)


def export_script_to_word(script_data: list, output_path: Path) -> None:
    """
    将剧本 JSON 导出为 Word 文档。

    只包含：场景标题、剧情描述（what）、对话（speaker + content）、镜头描述（shot_description）。
    跳过所有技术字段：shot_type, shot_blend, Follow, actions, current position, initial position, move。
    """
    doc = Document()

    # 文档标题
    title_para = doc.add_heading("剧 本", level=0)
    title_para.alignment = 1  # 居中
    for run in title_para.runs:
        _set_chinese_font(run)

    for i, scene_obj in enumerate(script_data):
        if not isinstance(scene_obj, dict):
            continue

        info = scene_obj.get("scene information", {})
        where = info.get("where", "")
        what = info.get("what", "")

        # 幕标题
        act_heading = doc.add_heading(f"第 {i + 1} 幕  {where}", level=1)
        for run in act_heading.runs:
            _set_chinese_font(run)

        # 剧情描述
        if what:
            p = doc.add_paragraph()
            run = p.add_run(f"[剧情] {what}")
            run.bold = True
            run.font.size = Pt(11)
            _set_chinese_font(run)
            p.paragraph_format.space_after = Pt(6)

        # 对话与镜头描述
        for seg in scene_obj.get("scene", []):
            if not isinstance(seg, dict):
                continue

            # 跳过纯移动片段（无 speaker/content）
            if "move" in seg and "speaker" not in seg:
                continue

            speaker = seg.get("speaker", "")
            content = seg.get("content", "")
            shot_desc = seg.get("shot_description", "")
            is_empty_shot = "speaker" in seg and "content" in seg and not str(speaker).strip() and not str(content).strip()

            if not speaker and not content and not is_empty_shot:
                continue

            # 对话行
            p = doc.add_paragraph()
            if is_empty_shot:
                duration = seg.get("duration") or "5s"
                empty_run = p.add_run(f"[空镜] {duration}")
                empty_run.bold = True
                empty_run.font.size = Pt(10.5)
                empty_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
                _set_chinese_font(empty_run)
            elif speaker:
                name_run = p.add_run(f"{speaker}：")
                name_run.bold = True
                name_run.font.size = Pt(10.5)
                _set_chinese_font(name_run)
            if content:
                text_run = p.add_run(content)
                text_run.font.size = Pt(10.5)
                _set_chinese_font(text_run)
            p.paragraph_format.space_after = Pt(2)

            # 镜头描述（缩进、斜体、灰色）
            if shot_desc:
                pd = doc.add_paragraph()
                pd.paragraph_format.left_indent = Cm(0.8)
                pd.paragraph_format.space_after = Pt(6)
                desc_run = pd.add_run(f"[镜头] {shot_desc}")
                desc_run.italic = True
                desc_run.font.size = Pt(9.5)
                desc_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
                _set_chinese_font(desc_run)

        # 幕间分隔
        sep = doc.add_paragraph()
        sep.paragraph_format.space_after = Pt(12)

    doc.save(str(output_path))
